"""docx → 공통 구조 파서 (포맷별 어댑터).

Word 문서를 '평탄화'하지 않고, 읽기 순서대로 문단/표를 위치 anchor와 함께 표현한다.
- anchor: 나중에 그 자리에 정확히 내용을 주입(in-place fill)하기 위한 좌표
    문단: {"type": "para", "p": <document.paragraphs 인덱스>}
    셀:   {"type": "cell", "t": <표 인덱스>, "r": <행>, "c": <열>}
    표:   {"type": "table", "t": <표 인덱스>}
- 빈칸 판정(fillable/fill_mode/label)은 templates.blanks.classify_blank 에 위임한다.
"""
import io

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.templates.blanks import classify_blank, classify_cell_blank


def _iter_body_blocks(doc):
    """문서 본문을 '읽기 순서'대로 (Paragraph | Table) 로 순회."""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _is_heading(para: Paragraph) -> bool:
    name = (para.style.name if para.style else "") or ""
    return name.startswith("Heading") or name == "Title"


def parse_structure(data: bytes) -> dict:
    """docx 바이트 → 공통 구조 dict. {"elements": [...]}.

    문단: {id, kind: paragraph|heading, style, text, empty, fillable, fill_mode, label, anchor}
    표:   {id, kind: table, rows, cols, cells:[{id, text, empty, fillable, fill_mode, anchor}], anchor}
    """
    doc = Document(io.BytesIO(data))
    para_index = {p._p: i for i, p in enumerate(doc.paragraphs)}

    elements: list[dict] = []
    eid = 0
    t_idx = -1
    for block in _iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text
            heading = _is_heading(block)
            fillable, fill_mode, label = classify_blank(text, is_heading=heading)
            elements.append(
                {
                    "id": eid,
                    "kind": "heading" if heading else "paragraph",
                    "style": block.style.name if block.style else "",
                    "text": text,
                    "empty": not text.strip(),
                    "fillable": fillable,
                    "fill_mode": fill_mode,
                    "label": label,
                    "anchor": {"type": "para", "p": para_index.get(block._p, -1)},
                }
            )
            eid += 1
        elif isinstance(block, Table):
            t_idx += 1
            cells: list[dict] = []
            for r, row in enumerate(block.rows):
                # python-docx는 가로 병합 셀을 grid 폭만큼 같은 _tc로 반복 반환한다.
                # 중복 슬롯은 fillable로 세지 않아 방향 감지·주입이 왜곡되지 않게 한다.
                seen_tc: set[int] = set()
                for c, cell in enumerate(row.cells):
                    ctext = cell.text
                    empty = not ctext.strip()
                    tc_id = id(cell._tc)
                    merged_skip = tc_id in seen_tc
                    seen_tc.add(tc_id)
                    if merged_skip:
                        fillable, fill_mode, label = False, "", ""
                    else:
                        fillable, fill_mode, label = classify_cell_blank(ctext)
                    # 셀은 표(top-level) 묶음의 일부 — 별도 id 없이 anchor만 (채우기 주입용)
                    cells.append(
                        {
                            "kind": "cell",
                            "text": ctext,
                            "empty": empty,
                            "fillable": fillable,
                            "fill_mode": fill_mode or ("replace" if fillable else ""),
                            "label": label,
                            "merged_skip": merged_skip,
                            "anchor": {"type": "cell", "t": t_idx, "r": r, "c": c},
                        }
                    )
            elements.append(
                {
                    "id": eid,  # 표도 하나의 top-level 요소 → 묶음(segment) 참조 단위
                    "kind": "table",
                    "rows": len(block.rows),
                    "cols": len(block.rows[0].cells) if block.rows else 0,
                    "cells": cells,
                    "anchor": {"type": "table", "t": t_idx},
                }
            )
            eid += 1
    return {"elements": elements}
