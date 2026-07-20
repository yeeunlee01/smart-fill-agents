"""채운 내용을 템플릿 docx의 anchor 위치에 in-place 주입 → 채워진 docx.

- 텍스트 slot(제목+빈 문단 등): 빈 문단에 내용 주입 (줄바꿈 유지). "라벨:"이면 라벨 뒤에 이어붙임.
- 목록 slot(불릿/번호): 예시 항목 문단을 틀로 복제해 항목 수만큼 채움 (표 행 확장과 동일).
- 표 slot: 템플릿의 '빈 행'을 반복 틀로 삼아, 내용(마크다운 표)의 행 수만큼 복제해 채운다.

주의: 일반 문단은 텍스트만 바꿔 인덱스를 유지. 목록/표 복제 시 뒤쪽 문단 인덱스가
      밀릴 수 있어 injections는 문서 뒤쪽(element id 큰 순)부터 적용한다.
"""
import copy
import io
import re

from docx import Document
from docx.text.paragraph import Paragraph

from app.templates.parsers.docx import parse_structure
from app.templates.table_layout import extract_kv_pairs, is_text_box_table, normalize_label, orientation

# blanks._LIST_ITEM 과 동일 계열 — 주입 시 마크다운 목록 줄 파싱용
_LIST_LINE = re.compile(
    r"^(?P<marker>[-*•·○●◦▪▸►]|\d+[.)]|[①-⑳]|[가나다라마바사아자차카타파하][.)]|[a-zA-Z][.)])\s*(?P<body>.*)$"
)


def fill_docx(data: bytes, injections: list[dict]) -> bytes:
    """injections: [{"element_ids": [int], "content": str, "repeatable": bool}] → 채워진 docx bytes."""
    doc = Document(io.BytesIO(data))
    structure = parse_structure(data)
    by_id = {el["id"]: el for el in structure["elements"]}

    # 목록/표 복제가 뒤 문단 인덱스를 밀지 않도록 문서 뒤에서부터 주입
    ordered = sorted(
        injections,
        key=lambda inj: max(inj.get("element_ids") or [-1]),
        reverse=True,
    )
    for inj in ordered:
        els = [by_id[i] for i in (inj.get("element_ids") or []) if i in by_id]
        content = (inj.get("content") or "").strip()
        if not els or not content or _is_not_found(content):
            continue
        table_el = next((e for e in els if e["kind"] == "table"), None)
        if table_el and is_text_box_table(table_el["cells"], table_el["rows"], table_el["cols"]):
            # Word 네모 작성란(1칸 표) → 줄글만 셀에 넣음 (마크다운 표로 넣지 않음)
            _inject_text_box(doc, table_el, content)
        elif table_el:
            # repeatable=False면 고정 개수 표(결재 등) → 행/열 복제 금지, 있는 빈칸만 채움. (기본 True: 기존 동작)
            _inject_table(doc, table_el, content, repeatable=inj.get("repeatable", True))
        elif any(e.get("fill_mode") == "list_item" for e in els):
            _inject_list(doc, els, content, repeatable=inj.get("repeatable", True))
        else:
            _inject_text(doc, els, content)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _is_not_found(content: str) -> bool:
    return "찾지 못했" in content or "찾을 수 없" in content


# ── 텍스트 slot ─────────────────────────────────────────────
def _clean_line(s: str) -> str:
    """마크다운 표기를 문서용 평문으로 가볍게 정리."""
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)          # **굵게** → 굵게
    s = re.sub(r"^\s*[-*]\s+", "• ", s)             # - 목록 → • 목록
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s.rstrip()


def _set_paragraph_text(para, text: str):
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    para.add_run(text)


def _clone_paragraph_after(para: Paragraph) -> Paragraph:
    """목록 틀 문단을 바로 뒤에 복제."""
    new_p = copy.deepcopy(para._p)
    para._p.addnext(new_p)
    return Paragraph(new_p, para._parent)


def _parse_list_items(content: str) -> list[str]:
    """LLM 출력 → 목록 본문들. `- 항목` / `• 항목` / 일반 줄 모두 허용."""
    items: list[str] = []
    for line in content.split("\n"):
        s = line.strip()
        if not s:
            continue
        m = _LIST_LINE.match(s)
        if m:
            body = (m.group("body") or "").strip()
            if body:
                items.append(body)
        else:
            items.append(s)
    return items


def _inject_list(doc, els: list[dict], content: str, repeatable: bool = True):
    """불릿/번호 틀 문단을 항목 수만큼 복제해 채운다."""
    list_els = [
        e for e in els
        if e.get("fill_mode") == "list_item" and e.get("anchor", {}).get("type") == "para"
    ]
    if not list_els:
        _inject_text(doc, els, content)
        return
    items = _parse_list_items(content)
    if not items:
        return
    if not repeatable:
        items = items[: len(list_els)]

    marker = (list_els[0].get("label") or "•").strip() or "•"
    paras: list[Paragraph] = []
    for e in list_els:
        p_idx = e["anchor"]["p"]
        if 0 <= p_idx < len(doc.paragraphs):
            paras.append(doc.paragraphs[p_idx])
    if not paras:
        return

    last = paras[-1]
    while len(paras) < len(items):
        last = _clone_paragraph_after(last)
        paras.append(last)

    for para, item in zip(paras, items):
        # 숫자 마커면 항목마다 번호를 유지하기보다 템플릿 마커를 그대로 씀 (• 틀이 일반적)
        prefix = f"{marker} " if marker else ""
        _set_paragraph_text(para, f"{prefix}{item}".rstrip())


def _inject_text(doc, els: list[dict], content: str):
    # slot 안의 첫 fillable 문단을 주입 지점으로 (없으면 마지막 요소)
    target = next((e for e in els if e.get("fillable") and e["anchor"]["type"] == "para"), None)
    if target is None:
        return
    p_idx = target["anchor"]["p"]
    if p_idx < 0 or p_idx >= len(doc.paragraphs):
        return
    para = doc.paragraphs[p_idx]

    lines = [_clean_line(x) for x in content.split("\n")]
    lines = [x for x in lines if x != ""] or [content]

    if target.get("fill_mode") == "append":
        # "라벨: " → 기존 텍스트(라벨) 뒤에 값을 이어붙임
        para.add_run(" " + " ".join(lines))
        return
    # 빈 문단 → 기존 run 제거 후 내용으로 채움 (여러 줄은 줄바꿈으로)
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for i, ln in enumerate(lines):
        if i > 0:
            para.add_run().add_break()
        para.add_run(ln)


# ── 표 slot ─────────────────────────────────────────────────
_SEP = re.compile(r"[\s:|-]+")


def _inject_text_box(doc, table_el: dict, content: str):
    """1칸 작성란 표: 셀 전체를 줄글(여러 문단)로 교체. 마크다운 표 기호는 평문으로 정리.

    빈 줄로 문단을 나누고, 문단 안 줄바꿈은 한 문단으로 이어 붙인다.
    """
    t_idx = table_el["anchor"]["t"]
    if t_idx < 0 or t_idx >= len(doc.tables):
        return
    cell = doc.tables[t_idx].rows[0].cells[0]

    def _clean_line_keep(s: str) -> str:
        if _SEP.fullmatch(s) and "-" in s and "|" in s:
            return ""
        if "|" in s:
            parts = re.sub(r"^\s*\|", "", re.sub(r"\|\s*$", "", s)).split("|")
            return " ".join(_clean_line(p) for p in parts if p.strip())
        return _clean_line(s)

    # 빈 줄 기준 문단 분리
    paragraphs: list[str] = []
    buf: list[str] = []
    for line in content.split("\n"):
        s = line.strip()
        if not s:
            if buf:
                paragraphs.append(" ".join(buf))
                buf = []
            continue
        cleaned = _clean_line_keep(s)
        if cleaned:
            buf.append(cleaned)
    if buf:
        paragraphs.append(" ".join(buf))
    if not paragraphs:
        return
    for p in cell.paragraphs:
        for rn in list(p.runs):
            rn._element.getparent().remove(rn._element)
    cell.paragraphs[0].add_run(paragraphs[0])
    for para in paragraphs[1:]:
        p = cell.add_paragraph()
        p.add_run(para)


def _parse_grid(content: str):
    """내용 → (2D 그리드, 마크다운표였나). 표면 각 행이 셀 리스트, 아니면 각 줄이 1열."""
    grid = []
    is_table = False
    for line in content.split("\n"):
        s = line.strip()
        if not s:
            continue
        if "|" in s:
            if _SEP.fullmatch(s) and "-" in s:  # |---|---| 구분행 → 표 확정
                is_table = True
                continue
            cells = re.sub(r"^\s*\|", "", re.sub(r"\|\s*$", "", s)).split("|")
            grid.append([_clean_line(c) for c in cells])
        else:
            grid.append([_clean_line(s)])
    return grid, is_table


def _set_cell_text(cell, val: str):
    for p in cell.paragraphs:
        for rn in list(p.runs):
            rn._element.getparent().remove(rn._element)
    cell.paragraphs[0].add_run(val)


def _write_cell(cell, val: str, fill_mode: str = "replace", hint: str = ""):
    """fill_mode에 맞게 셀에 값 기록. prefix는 단위/안내(hint)를 값 뒤에 유지."""
    if fill_mode == "prefix":
        suffix = (hint or cell.text or "").strip()
        _set_cell_text(cell, f"{val} {suffix}".strip() if suffix else val)
    elif fill_mode == "append":
        base = (cell.text or "").rstrip()
        _set_cell_text(cell, f"{base} {val}".strip() if base else val)
    else:
        _set_cell_text(cell, val)


def _clone_row_after(table, ref_row):
    """빈 행(반복 틀)을 바로 뒤에 복제해 새 행을 반환 (아래로 확장)."""
    import copy

    from docx.table import _Row
    new_tr = copy.deepcopy(ref_row._tr)
    ref_row._tr.addnext(new_tr)
    return _Row(new_tr, table)


def _clone_col_after(table, col_idx: int) -> int:
    """col_idx 열을 오른쪽에 복제(오른쪽으로 확장). 새 열 인덱스 반환. (병합 열은 미지원)"""
    import copy

    from docx.oxml.ns import qn
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        gcs = grid.findall(qn("w:gridCol"))
        if 0 <= col_idx < len(gcs):
            gcs[col_idx].addnext(copy.deepcopy(gcs[col_idx]))
    for tr in tbl.findall(qn("w:tr")):
        tcs = tr.findall(qn("w:tc"))
        if 0 <= col_idx < len(tcs):
            tcs[col_idx].addnext(copy.deepcopy(tcs[col_idx]))
    return col_idx + 1


def _inject_table(doc, table_el: dict, content: str, repeatable: bool = True):
    """열/행 고정 + 나머지 축 동적. 방향 감지 후 행 복제(아래로) 또는 열 복제(오른쪽)로 채운다.
    양식(kv) 표는 라벨→값 셀을 항목명으로 매핑해 채운다 (한 행에 쌍이 여러 개여도 가능).
    헤더·병합·라벨 셀(=fillable 아님)은 안 건드림. 마크다운 표 헤더행은 제외.
    repeatable=False면 고정 개수 표(결재 등) → 복제 없이 있는 빈칸만 채운다.
    """
    t_idx = table_el["anchor"]["t"]
    if t_idx < 0 or t_idx >= len(doc.tables):
        return
    table = doc.tables[t_idx]
    rows, cols = table_el["rows"], table_el["cols"]
    cells = table_el["cells"]
    grid, is_table = _parse_grid(content)
    if is_table and grid:
        grid = grid[1:]  # 헤더행(필드 이름) 제외 — 이름은 템플릿 것 사용
    grid = [g for g in grid if any(v for v in g)]
    if not grid:
        return

    axis = orientation(cells, rows, cols)
    if axis == "kv":
        _inject_kv(table, cells, rows, cols, grid)
    elif axis == "row":
        _inject_rowwise(table, cells, rows, cols, grid, repeatable)
    else:
        _inject_colwise(table, cells, rows, cols, grid, repeatable)


def _inject_kv(table, cells, rows, cols, grid):
    """양식 표: | 항목 | 값 | 행들을 라벨로 매칭해 해당 값 셀에 쓴다."""
    pairs = extract_kv_pairs(cells, rows, cols)
    if not pairs:
        return
    by_label = {normalize_label(p["label"]): p for p in pairs}
    used: set[tuple[int, int]] = set()

    def _fill_pair(p: dict, val: str):
        key = (p["r"], p["c"])
        if key in used or not val:
            return
        used.add(key)
        hint = cells[p["r"] * cols + p["c"]].get("text") or ""
        _write_cell(table.rows[p["r"]].cells[p["c"]], val, p.get("fill_mode") or "replace", hint)

    for rec in grid:
        if len(rec) < 2 or not (rec[1] or "").strip():
            continue
        label, val = (rec[0] or "").strip(), (rec[1] or "").strip()
        p = by_label.get(normalize_label(label))
        if p is not None:
            _fill_pair(p, val)
    # 라벨 매칭 실패 시 남은 쌍에 순서대로 값만 있는 행을 보충
    leftovers = [(r[0] or "").strip() for r in grid if len(r) == 1 and (r[0] or "").strip()]
    unset = [p for p in pairs if (p["r"], p["c"]) not in used]
    for p, val in zip(unset, leftovers):
        _fill_pair(p, val)


def _inject_rowwise(table, cells, rows, cols, grid, repeatable: bool = True):
    """행 방향: 각 레코드=한 행. 빈 행에 채우고, 레코드가 더 많으면 빈 행 복제(아래로)."""
    blank_rows = [
        r for r in range(rows)
        if any(
            cells[r * cols + c].get("fillable") and not cells[r * cols + c].get("merged_skip")
            for c in range(cols)
        )
    ]
    if not blank_rows:
        return
    if not repeatable:
        grid = grid[:len(blank_rows)]  # 고정 개수 표: 빈 행만큼만, 복제 없음
    blank_cols = [
        c for c in range(cols)
        if cells[blank_rows[0] * cols + c].get("fillable")
        and not cells[blank_rows[0] * cols + c].get("merged_skip")
    ]
    last_row = None
    for i, rec in enumerate(grid):
        row = table.rows[blank_rows[i]] if i < len(blank_rows) else _clone_row_after(table, last_row)
        last_row = row                             # 다음 복제는 '직전 행' 뒤에 → 순서 유지
        for c in blank_cols:                        # 절대 열 매핑: rec[c] → 열 c
            if c < len(rec) and rec[c]:
                cell_meta = cells[blank_rows[min(i, len(blank_rows) - 1)] * cols + c] if i < len(blank_rows) else {}
                _write_cell(row.cells[c], rec[c], cell_meta.get("fill_mode") or "replace", cell_meta.get("text") or "")


def _inject_colwise(table, cells, rows, cols, grid, repeatable: bool = True):
    """열 방향: 각 시리즈=한 열. 빈 열에 채우고, 시리즈가 더 많으면 빈 열 복제(오른쪽).

    grid[r] = [라벨, 값1, 값2, ...]. 라벨(0열)은 템플릿에 있으니 제외, 값 열들이 시리즈.
    """
    blank_cols = [
        c for c in range(cols)
        if any(
            cells[r * cols + c].get("fillable") and not cells[r * cols + c].get("merged_skip")
            for r in range(rows)
        )
    ]
    if not blank_cols:
        return
    blank_rows = [
        r for r in range(rows)
        if cells[r * cols + blank_cols[0]].get("fillable")
        and not cells[r * cols + blank_cols[0]].get("merged_skip")
    ]
    n_series = max((len(g) for g in grid), default=1) - 1  # 라벨 열 제외
    n_series = max(n_series, 1)
    if not repeatable:
        n_series = min(n_series, len(blank_cols))  # 고정 개수 표: 빈 열만큼만, 복제 없음
    for k in range(n_series):                              # k번째 시리즈 → grid의 (k+1)열
        tcol = blank_cols[k] if k < len(blank_cols) else _clone_col_after(table, blank_cols[-1])
        for i, r in enumerate(blank_rows):                 # 행은 라벨 순서대로 정렬
            g = grid[i] if i < len(grid) else []
            val = g[k + 1] if (k + 1) < len(g) else (g[-1] if len(g) == 1 else "")
            if val:
                cell_meta = cells[r * cols + tcol] if tcol < cols else {}
                _write_cell(
                    table.rows[r].cells[tcol], val,
                    cell_meta.get("fill_mode") or "replace", cell_meta.get("text") or "",
                )
