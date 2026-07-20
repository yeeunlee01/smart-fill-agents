"""양식(kv) 표: 한 행에 라벨→값 쌍이 2개인 경우 인식·주입."""
import io
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[2]
BE = ROOT / "be"
SAMPLE = ROOT / "data" / "시연" / "사업제안요약서.docx"

pytest.importorskip("docx")
import sys

sys.path.insert(0, str(BE))

from app.templates.blanks import classify_cell_blank  # noqa: E402
from app.templates.inject import fill_docx  # noqa: E402
from app.templates.parsers.docx import parse_structure  # noqa: E402
from app.templates.table_layout import orientation, table_fields  # noqa: E402


@pytest.mark.skipif(not SAMPLE.exists(), reason="sample docx missing")
def test_dual_pair_form_tables_are_kv():
    s = parse_structure(SAMPLE.read_bytes())
    tables = [e for e in s["elements"] if e["kind"] == "table"]
    overview, vendor = tables[0], tables[1]

    assert orientation(overview["cells"], overview["rows"], overview["cols"]) == "kv"
    fields = table_fields(overview["cells"], overview["rows"], overview["cols"])
    assert "사 업 기 간" in fields
    assert "제안서 제출 마감일" in fields
    assert "총 제안금액" in fields

    assert orientation(vendor["cells"], vendor["rows"], vendor["cols"]) == "kv"
    vfields = table_fields(vendor["cells"], vendor["rows"], vendor["cols"])
    assert "대 표 자" in vfields
    assert "연 락 처" in vfields


@pytest.mark.skipif(not SAMPLE.exists(), reason="sample docx missing")
def test_record_tables_stay_row():
    s = parse_structure(SAMPLE.read_bytes())
    tables = [e for e in s["elements"] if e["kind"] == "table"]
    # 주요 수행실적
    perf = tables[3]
    assert orientation(perf["cells"], perf["rows"], perf["cols"]) == "row"


def test_unit_hint_not_confused_with_labels():
    assert classify_cell_blank("대 표 자") == (False, "", "")
    assert classify_cell_blank("담 당 자") == (False, "", "")
    fillable, mode, _ = classify_cell_blank("원  ( 부가가치세 별도 )")
    assert fillable and mode == "prefix"


@pytest.mark.skipif(not SAMPLE.exists(), reason="sample docx missing")
def test_inject_dual_pair_and_prefix():
    from docx import Document

    data = SAMPLE.read_bytes()
    s = parse_structure(data)
    t0 = next(e for e in s["elements"] if e["kind"] == "table")
    out = fill_docx(
        data,
        [{
            "element_ids": [t0["id"]],
            "repeatable": False,
            "content": (
                "| 항목 | 값 |\n| --- | --- |\n"
                "| 사 업 명 | A |\n"
                "| 발 주 기 관 | B |\n"
                "| 제안서 제출 마감일 | D1 |\n"
                "| 사 업 기 간 | P1 |\n"
                "| 총 제안금액 | 100 |\n"
            ),
        }],
    )
    doc = Document(io.BytesIO(out))
    row2 = doc.tables[0].rows[2]
    # 논리 셀만: 마감일 | D1 | 사업기간 | P1
    seen, texts = set(), []
    for cell in row2.cells:
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        texts.append(cell.text.strip())
    assert texts == ["제안서 제출 마감일", "D1", "사 업 기 간", "P1"]
    amount = doc.tables[0].rows[3].cells[1].text.strip()
    assert amount.startswith("100")
    assert "원" in amount
